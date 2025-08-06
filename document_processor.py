# app/services/document_processor.py - Document Processing Service
import aiohttp
import asyncio
from typing import List, Dict, Any, Optional, Union
import io
import re
from datetime import datetime
from dataclasses import dataclass
from pathlib import Path
import hashlib
import mimetypes

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import email
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import chardet

from loguru import logger
from settings import settings
from exceptions import DocumentProcessingError
from database import DatabaseManager

@dataclass
class DocumentContent:
    """Document content structure"""
    document_id: str
    title: Optional[str]
    content: str
    document_type: str
    source_url: Optional[str]
    metadata: Dict[str, Any]
    file_size: int
    created_at: datetime

@dataclass
class DocumentChunk:
    """Document chunk structure"""
    chunk_id: str
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any]

class DocumentProcessorService:
    """Service for processing various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/msword': self._process_doc,
            'text/plain': self._process_text,
            'message/rfc822': self._process_email,
            'text/html': self._process_html
        }
        self.max_file_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024  # Convert to bytes
    
    async def process_document_from_url(self, url: str) -> DocumentContent:
        """
        Download and process document from URL
        
        Args:
            url: Document URL to process
            
        Returns:
            DocumentContent: Processed document content
            
        Raises:
            DocumentProcessingError: If processing fails
        """
        try:
            logger.info(f"Processing document from URL: {url}")
            
            # Download document
            document_data, content_type, filename = await self._download_document(url)
            
            # Detect file type if not provided
            if not content_type:
                content_type = self._detect_file_type(document_data, filename)
            
            # Validate file size
            if len(document_data) > self.max_file_size:
                raise DocumentProcessingError(
                    f"File size ({len(document_data)} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)"
                )
            
            # Process based on content type
            processor = self.supported_formats.get(content_type)
            if not processor:
                raise DocumentProcessingError(f"Unsupported file type: {content_type}")
            
            # Extract content
            content, title, metadata = await processor(document_data, filename)
            
            # Create document ID
            document_id = self._generate_document_id(url, content)
            
            # Create document content object
            document_content = DocumentContent(
                document_id=document_id,
                title=title or self._extract_title_from_filename(filename),
                content=content,
                document_type=self._get_document_type(content_type),
                source_url=url,
                metadata={
                    **metadata,
                    'content_type': content_type,
                    'filename': filename,
                    'file_size': len(document_data),
                    'processed_at': datetime.utcnow().isoformat()
                },
                file_size=len(document_data),
                created_at=datetime.utcnow()
            )
            
            # Store in database
            await self._store_document(document_content)
            
            logger.info(f"Successfully processed document: {document_id}")
            return document_content
            
        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")
    
    async def _download_document(self, url: str) -> tuple[bytes, Optional[str], Optional[str]]:
        """Download document from URL"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=aiohttp.ClientTimeout(total=60)) as response:
                    if response.status != 200:
                        raise DocumentProcessingError(f"Failed to download document: HTTP {response.status}")
                    
                    content_type = response.headers.get('content-type')
                    content_disposition = response.headers.get('content-disposition', '')
                    
                    # Extract filename from URL or content-disposition
                    filename = self._extract_filename(url, content_disposition)
                    
                    # Read document data
                    document_data = await response.read()
                    
                    return document_data, content_type, filename
                    
        except aiohttp.ClientError as e:
            raise DocumentProcessingError(f"Failed to download document: {str(e)}")
    
    def _extract_filename(self, url: str, content_disposition: str) -> Optional[str]:
        """Extract filename from URL or content-disposition header"""
        # Try content-disposition header first
        if content_disposition:
            filename_match = re.search(r'filename="?([^"]+)"?', content_disposition)
            if filename_match:
                return filename_match.group(1)
        
        # Extract from URL
        try:
            path = Path(url.split('?')[0])  # Remove query parameters
            if path.suffix:
                return path.name
        except:
            pass
        
        return None
    
    def _detect_file_type(self, data: bytes, filename: Optional[str]) -> str:
        """Detect file type from data and filename"""
        # Try to detect from filename extension
        if filename:
            mime_type, _ = mimetypes.guess_type(filename)
            if mime_type:
                return mime_type
        
        # Magic number detection
        if data.startswith(b'%PDF'):
            return 'application/pdf'
        elif data.startswith(b'PK'):  # ZIP-based formats (DOCX)
            if b'word/' in data[:1000]:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif data.startswith(b'\xd0\xcf\x11\xe0'):  # OLE format (DOC)
            return 'application/msword'
        
        # Try to decode as text
        try:
            data.decode('utf-8')
            return 'text/plain'
        except UnicodeDecodeError:
            pass
        
        # Default to binary
        return 'application/octet-stream'
    
    async def _process_pdf(self, data: bytes, filename: Optional[str]) -> tuple[str, Optional[str], Dict[str, Any]]:
        """Process PDF document"""
        try:
            pdf_file = io.BytesIO(data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
            
            content = '\n\n'.join(text_content)
            
            # Extract metadata
            metadata = {
                'page_count': len(pdf_reader.pages),
                'format': 'PDF'
            }
            
            # Try to extract title from PDF metadata
            title = None
            if pdf_reader.metadata:
                title = pdf_reader.metadata.get('/Title')
                if pdf_reader.metadata.get('/Author'):
                    metadata['author'] = pdf_reader.metadata.get('/Author')
                if pdf_reader.metadata.get('/Subject'):
                    metadata['subject'] = pdf_reader.metadata.get('/Subject')
            
            if not content.strip():
                raise DocumentProcessingError("No text content found in PDF")
            
            return content, title, metadata
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process PDF: {str(e)}")
    
    async def _process_docx(self, data: bytes, filename: Optional[str]) -> tuple[str, Optional[str], Dict[str, Any]]:
        """Process DOCX document"""
        try:
            docx_file = io.BytesIO(data)
            document = DocxDocument(docx_file)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables_text = []
            for table in document.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            # Combine all text
            content_parts = paragraphs
            if tables_text:
                content_parts.extend(['', '--- Tables ---'] + tables_text)
            
            content = '\n'.join(content_parts)
            
            # Extract metadata
            metadata = {
                'format': 'DOCX',
                'paragraph_count': len(paragraphs),
                'table_count': len(document.tables)
            }
            
            # Try to extract title from document properties
            title = None
            if hasattr(document, 'core_properties'):
                title = document.core_properties.title
                if document.core_properties.author:
                    metadata['author'] = document.core_properties.author
                if document.core_properties.subject:
                    metadata['subject'] = document.core_properties.subject
            
            if not content.strip():
                raise DocumentProcessingError("No text content found in DOCX")
            
            return content, title, metadata
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process DOCX: {str(e)}")
    
    async def _process_doc(self, data: bytes, filename: Optional[str]) -> tuple[str, Optional[str], Dict[str, Any]]:
        """Process DOC document (legacy format)"""
        # For now, return error - would need python-docx2txt or similar
        raise DocumentProcessingError("DOC format not supported. Please convert to DOCX or PDF.")
    
    async def _process_text(self, data: bytes, filename: Optional[str]) -> tuple[str, Optional[str], Dict[str, Any]]:
        """Process plain text document"""
        try:
            # Detect encoding
            encoding_result = chardet.detect(data)
            encoding = encoding_result.get('encoding', 'utf-8')
            
            # Decode text
            content = data.decode(encoding)
            
            metadata = {
                'format': 'TEXT',
                'encoding': encoding,
                'confidence': encoding_result.get('confidence', 0.0)
            }
            
            title = self._extract_title_from_content(content)
            
            return content, title, metadata
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process text file: {str(e)}")
    
    async def _process_email(self, data: bytes, filename: Optional[str]) -> tuple[str, Optional[str], Dict[str, Any]]:
        """Process email document"""
        try:
            # Parse email
            msg = email.message_from_bytes(data)
            
            # Extract basic info
            subject = msg.get('Subject', '')
            sender = msg.get('From', '')
            recipient = msg.get('To', '')
            date = msg.get('Date', '')
            
            # Extract body
            body_parts = []
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == 'text/plain':
                        payload = part.get_payload(decode=True)
                        if payload:
                            try:
                                body_parts.append(payload.decode('utf-8'))
                            except UnicodeDecodeError:
                                body_parts.append(payload.decode('latin-1'))
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    try:
                        body_parts.append(payload.decode('utf-8'))
                    except UnicodeDecodeError:
                        body_parts.append(payload.decode('latin-1'))
            
            # Combine email content
            content_parts = [
                f"Subject: {subject}",
                f"From: {sender}",
                f"To: {recipient}",
                f"Date: {date}",
                "",
                "--- Email Body ---"
            ]
            content_parts.extend(body_parts)
            
            content = '\n'.join(content_parts)
            
            metadata = {
                'format': 'EMAIL',
                'subject': subject,
                'sender': sender,
                'recipient': recipient,
                'date': date
            }
            
            return content, subject, metadata
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process email: {str(e)}")
    
    async def _process_html(self, data: bytes, filename: Optional[str]) -> tuple[str, Optional[str], Dict[str, Any]]:
        """Process HTML document"""
        try:
            from bs4 import BeautifulSoup
            
            # Detect encoding
            encoding_result = chardet.detect(data)
            encoding = encoding_result.get('encoding', 'utf-8')
            
            # Parse HTML
            html_content = data.decode(encoding)
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Extract title
            title = None
            title_tag = soup.find('title')
            if title_tag:
                title = title_tag.string
            
            metadata = {
                'format': 'HTML',
                'encoding': encoding
            }
            
            return content, title, metadata
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process HTML: {str(e)}")
    
    def _generate_document_id(self, url: str, content: str) -> str:
        """Generate unique document ID"""
        combined = f"{url}:{content[:1000]}"
        return hashlib.sha256(combined.encode()).hexdigest()[:16]
    
    def _get_document_type(self, content_type: str) -> str:
        """Get simplified document type"""
        type_mapping = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'text',
            'message/rfc822': 'email',
            'text/html': 'html'
        }
        return type_mapping.get(content_type, 'unknown')
    
    def _extract_title_from_filename(self, filename: Optional[str]) -> Optional[str]:
        """Extract title from filename"""
        if not filename:
            return None
        
        # Remove extension and clean up
        title = Path(filename).stem
        title = re.sub(r'[_-]', ' ', title)
        title = ' '.join(word.capitalize() for word in title.split())
        
        return title if title else None
    
    def _extract_title_from_content(self, content: str) -> Optional[str]:
        """Extract title from document content"""
        lines = content.strip().split('\n')
        if lines:
            # Use first non-empty line as title
            first_line = lines[0].strip()
            if first_line and len(first_line) < 100:
                return first_line
        
        return None
    
    async def chunk_document(
        self,
        document: DocumentContent,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ) -> List[DocumentChunk]:
        """
        Chunk document content for processing
        
        Args:
            document: Document to chunk
            chunk_size: Size of each chunk (optional)
            chunk_overlap: Overlap between chunks (optional)
            
        Returns:
            List[DocumentChunk]: List of document chunks
        """
        chunk_size = chunk_size or settings.CHUNK_SIZE
        chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        
        try:
            content = document.content
            chunks = []
            
            # Simple text chunking by character count
            start = 0
            chunk_index = 0
            
            while start < len(content):
                end = min(start + chunk_size, len(content))
                
                # Try to break at sentence boundary
                if end < len(content):
                    # Look for sentence endings
                    sentence_end = self._find_sentence_boundary(content, end)
                    if sentence_end > start:
                        end = sentence_end
                
                chunk_content = content[start:end].strip()
                
                if chunk_content:
                    chunk_id = f"{document.document_id}_chunk_{chunk_index}"
                    
                    chunk = DocumentChunk(
                        chunk_id=chunk_id,
                        content=chunk_content,
                        chunk_index=chunk_index,
                        start_char=start,
                        end_char=end,
                        metadata={
                            'document_id': document.document_id,
                            'chunk_size': len(chunk_content),
                            'document_type': document.document_type
                        }
                    )
                    
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Move start position with overlap
                start = max(start + chunk_size - chunk_overlap, end)
                
                if start >= len(content):
                    break
            
            logger.info(f"Created {len(chunks)} chunks for document {document.document_id}")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to chunk document: {str(e)}")
            raise DocumentProcessingError(f"Failed to chunk document: {str(e)}")
    
    def _find_sentence_boundary(self, text: str, position: int) -> int:
        """Find appropriate sentence boundary near position"""
        # Look for sentence endings within reasonable range
        search_range = min(200, len(text) - position)
        
        for i in range(search_range):
            char = text[position + i]
            if char in '.!?':
                # Check if it's likely end of sentence (followed by space or newline)
                next_pos = position + i + 1
                if next_pos < len(text) and text[next_pos] in ' \n\t':
                    return position + i + 1
        
        # If no sentence boundary found, return original position
        return position
    
    async def _store_document(self, document: DocumentContent):
        """Store document in database"""
        try:
            await DatabaseManager.create_document(
                content=document.content,
                title=document.title,
                document_type=document.document_type,
                source_url=document.source_url,
                metadata=document.metadata
            )
        except Exception as e:
            logger.error(f"Failed to store document in database: {str(e)}")
            # Don't raise error - document processing can continue without DB storage
    
    async def extract_metadata(self, document: DocumentContent) -> Dict[str, Any]:
        """Extract additional metadata from document"""
        try:
            metadata = document.metadata.copy()
            
            # Add content statistics
            content = document.content
            metadata.update({
                'character_count': len(content),
                'word_count': len(content.split()),
                'line_count': len(content.split('\n')),
                'paragraph_count': len([p for p in content.split('\n\n') if p.strip()])
            })
            
            # Extract potential key terms (simple approach)
            words = re.findall(r'\b[A-Za-z]{3,}\b', content.lower())
            word_freq = {}
            for word in words:
                word_freq[word] = word_freq.get(word, 0) + 1
            
            # Get top 20 most frequent words
            top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]
            metadata['top_words'] = [word for word, freq in top_words]
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract metadata: {str(e)}")
            return document.metadata